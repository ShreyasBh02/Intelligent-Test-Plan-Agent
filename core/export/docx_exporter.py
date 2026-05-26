import os
import tempfile
from docx import Document

async def export_to_docx(markdown_content: str, jira_id: str) -> str:
    doc = Document()
    doc.add_heading(f"Test Plan: {jira_id}", 0)
    
    lines = markdown_content.split('\n')
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped == "":
            pass
        else:
            doc.add_paragraph(stripped)
            
    temp_dir = tempfile.gettempdir()
    output_path = os.path.join(temp_dir, f"TestPlan_{jira_id}.docx")
    doc.save(output_path)
    
    return output_path
